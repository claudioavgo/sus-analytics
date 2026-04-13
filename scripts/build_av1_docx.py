"""Gera o Documento de Arquitetura da AV1 em DOCX."""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parents[1] / "docs" / "arquitetura_dados.docx"


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_paragraph(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table


def build():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(
        "SUS Analytics: Documento de Arquitetura (AV1)", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = sub.add_run("Fundamentos de Big Data, primeira entrega (13/04)")
    srun.italic = True
    srun.font.size = Pt(12)

    team = doc.add_paragraph()
    team.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trun = team.add_run(
        "Equipe: Bruno Ribeiro, Cláudio Alves e Vinícius Ventura"
    )
    trun.font.size = Pt(11)

    doc.add_paragraph()

    add_heading(doc, "1. Introdução", 1)
    add_paragraph(
        doc,
        "Este documento apresenta a arquitetura de dados do projeto SUS Analytics, "
        "construído para a primeira entrega (AV1) da disciplina de Fundamentos de "
        "Big Data. O pipeline cobre dados de internações hospitalares do SUS no "
        "estado de São Paulo, entre janeiro de 2020 e dezembro de 2023, com foco "
        "no impacto das ondas do COVID-19 sobre a rede pública.",
    )

    add_heading(doc, "2. Pergunta de Pesquisa", 1)
    add_paragraph(
        doc,
        "Como as ondas do COVID-19 impactaram o volume de internações e a "
        "mortalidade hospitalar no SUS-SP entre 2020 e 2023?",
        italic=True,
    )
    add_paragraph(doc, "Perguntas secundárias:")
    for item in [
        "Como evoluiu, mês a mês, o volume de internações por COVID (CID B342) em São Paulo?",
        "A taxa de óbito hospitalar subiu durante os picos da pandemia?",
        "O perfil da internação (dias de permanência e uso de UTI) mudou de uma onda para outra?",
        "Quais grupos demográficos (sexo e faixa etária) foram mais afetados?",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "3. Fonte de Dados", 1)
    add_paragraph(
        doc,
        "A base utilizada é a SIH/SUS, na forma de AIH Reduzida, publicada pelo "
        "DATASUS. O recorte é único e intencional: apenas o estado de São Paulo, "
        "entre janeiro de 2020 e dezembro de 2023. São Paulo concentra a maior "
        "rede hospitalar do SUS e os maiores picos de internação por COVID no "
        "país, o que torna o recorte coeso para a pergunta de pesquisa sem "
        "inflar o pipeline com dados alheios ao problema.",
    )
    add_table(
        doc,
        ["Fonte", "Descrição", "Formato", "Cobertura"],
        [
            [
                "SIH/SUS (DATASUS)",
                "Autorizações de Internação Hospitalar, AIH Reduzida",
                ".dbc (DBF comprimido)",
                "SP, Jan/2020 a Dez/2023 (48 arquivos)",
            ]
        ],
    )

    add_heading(doc, "4. Diagrama do Pipeline de Dados", 1)
    add_paragraph(
        doc,
        "O projeto segue a arquitetura medalhão (bronze, silver e gold). Na AV1 "
        "estão implementadas as camadas bronze e silver. A camada gold, com as "
        "tabelas agregadas que respondem à pergunta de pesquisa, fica para a AV2.",
    )
    add_code_block(
        doc,
        "[DATASUS, FTP publico]\n"
        "         |\n"
        "         v\n"
        "  [Bronze]                                      <- implementado (AV1)\n"
        "  data/bronze/sihsus/SP/*.dbc   (48 arquivos)\n"
        "  data/bronze/sihsus_sp_raw.parquet  (consolidado, 113 campos)\n"
        "         |\n"
        "         | src/ingest_bronze.py\n"
        "         v\n"
        "  [Silver]                                      <- implementado (AV1)\n"
        "  data/silver/sihsus_sp.parquet\n"
        "  17 campos selecionados, tipados, limpos\n"
        "  + enriquecimento: is_covid, ano, mes, ano_mes, faixa_etaria\n"
        "  + normalizacao Min-Max das variaveis numericas\n"
        "         |\n"
        "         | src/transform_silver.py\n"
        "         v\n"
        "  [Visualizacoes exploratorias]                 <- implementado (AV1)\n"
        "  notebooks/02_silver_visualizacoes.ipynb\n"
        "         |\n"
        "         v\n"
        "  [Gold]                                        <- previsto (AV2)\n"
        "  Tabelas agregadas (series temporais, mortalidade por onda etc.)",
    )

    add_heading(doc, "5. Etapas do Pipeline", 1)

    add_heading(doc, "5.1 Ingestão (Bronze)", 2)
    add_paragraph(
        doc,
        "Script responsável: src/ingest_bronze.py. Ele lê os 48 arquivos .dbc "
        "baixados do DATASUS, descomprime cada um via datasus-dbc ou dbfread e "
        "junta tudo em um único Parquet bruto, preservando os 113 campos "
        "originais. Nada é transformado nesta etapa: sem tipagem, sem seleção, "
        "sem filtro. O encoding latin-1 é declarado explicitamente na leitura "
        "para evitar problemas com acentuação.",
    )

    add_heading(doc, "5.2 Armazenamento", 2)
    add_paragraph(
        doc,
        "Todas as camadas são salvas em Parquet com PyArrow e compressão "
        "snappy. O formato colunar e comprimido é adequado para consultas "
        "analíticas e mantém o tamanho dos arquivos baixo. A estrutura em "
        "disco segue o padrão abaixo:",
    )
    add_code_block(
        doc,
        "data/\n"
        "  bronze/\n"
        "    sihsus/SP/RDSP2001.dbc ... RDSP2312.dbc  # 48 .dbc brutos\n"
        "    sihsus_sp_raw.parquet                    # consolidado (113 campos)\n"
        "  silver/\n"
        "    sihsus_sp.parquet                        # limpo, tipado, enriquecido",
    )

    add_heading(doc, "5.3 Transformação (Silver)", 2)
    add_paragraph(
        doc,
        "Script responsável: src/transform_silver.py. Ele aplica as seguintes "
        "transformações, nesta ordem:",
    )
    for item in [
        "Seleção de 17 campos relevantes, dos 113 originais.",
        "Tipagem pelo clean_dtypes: datas viram datetime; numéricos viram Int64 ou float; strings passam por strip e upper.",
        "Descarte de qualquer linha com nulo em uma das 17 colunas selecionadas.",
        "Descarte de linhas duplicadas em todas as colunas.",
        "Filtro de período: permanecem apenas registros entre 2020 e 2023.",
        "Enriquecimento linha a linha: is_covid (CID B342), ano, mes, ano_mes e faixa_etaria.",
        "Normalização Min-Max das colunas numéricas (DIAS_PERM, UTI_MES_TO, IDADE, VAL_TOT) no intervalo [0, 1], preservando as colunas originais.",
        "Auditoria final pelo audit_quality, com asserções de invariante (sem nulos, sem duplicatas, período válido, ranges de normalização). Se alguma falha, o pipeline para e nada é gravado.",
    ]:
        doc.add_paragraph(item, style="List Number")

    add_heading(doc, "5.4 Campos Selecionados na Silver", 2)
    add_table(
        doc,
        ["Campo", "Tipo", "Descrição"],
        [
            ["ANO_CMPT", "str", "Ano de processamento"],
            ["MES_CMPT", "str", "Mês de processamento"],
            ["DT_INTER", "date", "Data de internação"],
            ["DT_SAIDA", "date", "Data de saída"],
            ["MUNIC_RES", "str", "Município de residência do paciente (cód. IBGE)"],
            ["MUNIC_MOV", "str", "Município do estabelecimento"],
            ["DIAG_PRINC", "str", "Diagnóstico principal (CID-10)"],
            ["MORTE", "int", "Óbito (0 = Não, 1 = Sim)"],
            ["DIAS_PERM", "int", "Dias de permanência"],
            ["UTI_MES_TO", "int", "Dias de UTI no mês"],
            ["CAR_INT", "str", "Caráter da internação (eletiva/urgência)"],
            ["COMPLEX", "str", "Nível de complexidade"],
            ["SEXO", "str", "Sexo do paciente"],
            ["IDADE", "int", "Idade do paciente"],
            ["COD_IDADE", "str", "Unidade de medida da idade (4 = anos)"],
            ["VAL_TOT", "float", "Valor total da AIH (R$)"],
            ["CNES", "str", "Código CNES do estabelecimento"],
        ],
    )

    add_heading(doc, "6. Tecnologias Utilizadas", 1)
    add_table(
        doc,
        ["Camada", "Tecnologia", "Por que foi escolhida"],
        [
            [
                "Ingestão",
                "datasus-dbc, dbfread",
                "Leem arquivos .dbc sem precisar de conversão externa",
            ],
            [
                "Processamento",
                "pandas",
                "Atende com folga o volume atual e o time já domina a API",
            ],
            [
                "Armazenamento",
                "Parquet via pyarrow",
                "Formato colunar e comprimido, adequado a consultas analíticas",
            ],
            [
                "Análise",
                "Jupyter, matplotlib e seaborn",
                "Ferramentas usuais para exploração e gráficos estáticos",
            ],
            [
                "Versionamento",
                "Git e GitHub",
                "Controle de versão e colaboração entre os três autores",
            ],
        ],
    )

    add_heading(doc, "6.1 Tecnologias Pagas para Refinamento", 2)
    add_paragraph(
        doc,
        "Em uma versão profissional, o pipeline poderia incorporar serviços "
        "gerenciados e plataformas pagas. Elas ficaram fora da AV1 por restrição "
        "de custo e de escopo, mas valem ser registradas como direção de "
        "evolução do projeto:",
    )
    add_table(
        doc,
        ["Tecnologia", "Papel no pipeline", "Motivo de não usar agora"],
        [
            [
                "AWS S3",
                "Armazenar bronze, silver e gold em storage durável e versionado",
                "Custo, o projeto roda na máquina do aluno",
            ],
            [
                "Apache Spark (Databricks ou EMR)",
                "Processar a silver de forma distribuída em escala nacional",
                "O volume atual (48 arquivos, só SP) não justifica",
            ],
            [
                "dbt Cloud",
                "Transformações gold versionadas, testadas e documentadas",
                "Complexidade fora do escopo da AV1",
            ],
            [
                "Apache Airflow (MWAA)",
                "Orquestração ponta a ponta com agendamento periódico",
                "O pipeline ainda é disparado manualmente na AV1",
            ],
            [
                "Databricks Lakehouse",
                "Plataforma única para bronze, silver, gold e BI",
                "Custo e complexidade acima do necessário",
            ],
            [
                "Snowflake ou BigQuery",
                "Data warehouse gerenciado para a camada gold",
                "Parquet local atende ao volume atual",
            ],
        ],
    )

    add_heading(doc, "7. Qualidade dos Dados", 1)
    add_table(
        doc,
        ["Problema", "Descrição", "Tratamento"],
        [
            [
                "Campos legados",
                "19 campos zerados desde 2015",
                "Descartados já na seleção da silver",
            ],
            [
                "Nulos",
                "Qualquer coluna nula entre as 17 selecionadas",
                "Linha descartada por drop_nulls e validada por audit_quality",
            ],
            [
                "Duplicatas",
                "Linhas idênticas em todas as colunas",
                "Removidas por drop_duplicates",
            ],
            [
                "Encoding",
                "Arquivos em latin-1",
                "Encoding declarado explicitamente na leitura",
            ],
            [
                "CID do COVID",
                "O DATASUS usa B342, e não U071 ou U072",
                "A flag is_covid usa B342",
            ],
            [
                "Escalas",
                "Variáveis numéricas em unidades diferentes",
                "Colunas *_norm Min-Max em [0, 1]",
            ],
            [
                "Invariantes",
                "Nulos, duplicatas, período e intervalos de normalização",
                "audit_quality roda as asserções no fim do pipeline",
            ],
        ],
    )

    add_heading(doc, "8. Arquitetura Parcial Implementada (AV1)", 1)
    add_paragraph(
        doc,
        "Na AV1 estão prontas as camadas bronze e silver, além de um conjunto "
        "inicial de visualizações exploratórias sobre a silver. A camada gold, "
        "com as tabelas agregadas que respondem à pergunta de pesquisa, é "
        "escopo da AV2.",
    )
    add_table(
        doc,
        ["Componente", "Status AV1", "Artefato"],
        [
            ["Ingestão bronze", "Finalizado", "src/ingest_bronze.py"],
            [
                "Armazenamento bronze e silver",
                "Finalizado",
                "data/bronze, data/silver (Parquet)",
            ],
            ["Transformação silver", "Finalizado", "src/transform_silver.py"],
            [
                "Auditoria de qualidade",
                "Finalizado",
                "audit_quality em transform_silver.py",
            ],
            [
                "Demonstração técnica",
                "Finalizado",
                "notebooks/01_exploracao_sihsus.ipynb",
            ],
            [
                "Visualizações exploratórias",
                "Finalizado",
                "notebooks/02_silver_visualizacoes.ipynb",
            ],
            ["Camada gold (agregações)", "Previsto", "AV2"],
            ["Dashboard final", "Previsto", "AV2"],
        ],
    )

    add_heading(doc, "9. Equipe Responsável e Divisão de Tarefas", 1)
    add_table(
        doc,
        ["Membro", "Responsabilidades principais"],
        [
            [
                "Bruno Ribeiro",
                "Pipeline bronze para silver, scripts em Python e decisões de arquitetura",
            ],
            [
                "Vinícius Ventura",
                "Notebook de análise, visualizações exploratórias e apoio na redação",
            ],
            [
                "Cláudio Alves",
                "README, documentação técnica, checklist e preparação da apresentação em aula",
            ],
        ],
    )

    add_heading(doc, "9.1 Divisão de tarefas na AV1", 2)
    add_table(
        doc,
        ["Tarefa", "Responsável", "Status"],
        [
            [
                "Download e organização dos dados .dbc",
                "Bruno Ribeiro e Vinícius Ventura",
                "Concluído",
            ],
            [
                "Implementação do src/ingest_bronze.py",
                "Bruno Ribeiro e Vinícius Ventura",
                "Concluído",
            ],
            [
                "Implementação do src/transform_silver.py",
                "Bruno Ribeiro e Vinícius Ventura",
                "Concluído",
            ],
            [
                "Notebook de demonstração técnica",
                "Bruno Ribeiro e Vinícius Ventura",
                "Concluído",
            ],
            ["Documentação de arquitetura", "Cláudio Alves", "Concluído"],
            ["Dicionário de dados", "Cláudio Alves", "Concluído"],
            ["README e checklist da AV1", "Cláudio Alves", "Concluído"],
        ],
    )

    add_heading(doc, "10. Checklist AV1: Estado do Pipeline", 1)
    add_paragraph(
        doc,
        "Formulário oficial pedido na seção 6.1 da especificação da disciplina:",
    )
    add_code_block(
        doc,
        "Ingestão:      ( ) Em progresso  (x) Finalizado  ( ) Pendente\n"
        "Armazenamento: ( ) Em progresso  (x) Finalizado  ( ) Pendente\n"
        "Transformação: ( ) Em progresso  (x) Finalizado  ( ) Pendente",
    )
    add_paragraph(
        doc,
        "As três etapas obrigatórias do pipeline estão finalizadas para a AV1. "
        "A camada gold, com as tabelas agregadas, e o dashboard final ficam "
        "para a AV2.",
    )

    add_heading(doc, "11. Repositório GitHub", 1)
    add_paragraph(
        doc,
        "O código, a documentação e os notebooks estão versionados em "
        "repositório Git público, com commits visíveis dos três autores. A "
        "estrutura em disco segue o padrão abaixo:",
    )
    add_code_block(
        doc,
        "sus-analytics/\n"
        "  data/\n"
        "    bronze/   # 48 .dbc do DATASUS e parquet consolidado\n"
        "    silver/   # parquet limpo e enriquecido\n"
        "  src/\n"
        "    ingest_bronze.py       # .dbc -> parquet bruto\n"
        "    transform_silver.py    # bronze -> silver\n"
        "  notebooks/\n"
        "    01_exploracao_sihsus.ipynb\n"
        "    02_silver_visualizacoes.ipynb\n"
        "  docs/\n"
        "    arquitetura_dados.md\n"
        "    arquitetura_dados.docx   # este documento\n"
        "    checklist_av1.md\n"
        "    team_roles.md\n"
        "    PROJETO_FUND_BIG_DATA.pdf\n"
        "  README.md\n"
        "  requirements.txt",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"DOCX gerado em: {OUTPUT}")


if __name__ == "__main__":
    build()
